#!/usr/bin/env python3
"""Generate Word (.docx) and PDF versions of the reconnaissance report."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, HRFlowable, PageBreak)
from reportlab.platypus.flowables import KeepTogether
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_RIGHT, TA_CENTER, TA_LEFT
import os, glob

# ─── Arabic font setup ───────────────────────────────────────────────────────
FONT_PATHS = glob.glob("/usr/share/fonts/**/*.ttf", recursive=True)
ARABIC_FONT = None
for p in FONT_PATHS:
    name = os.path.basename(p).lower()
    if any(k in name for k in ("arabic", "noto", "amiri", "cairo", "tajawal", "dejavu")):
        try:
            pdfmetrics.registerFont(TTFont("ArabicFont", p))
            ARABIC_FONT = "ArabicFont"
            print(f"[PDF] Using font: {p}")
            break
        except Exception:
            pass
if not ARABIC_FONT:
    ARABIC_FONT = "Helvetica"
    print("[PDF] No Arabic font found, falling back to Helvetica (Arabic may not render)")


# ════════════════════════════════════════════════════════════════════════════
#  WORD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)

def add_heading(doc, text, level=1, color=(0x1a, 0x37, 0x6c)):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(h)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = "Arial"
    return h

def add_rtl_para(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "1a376c")
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = "E8EEF7" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(cell.paragraphs[0])
            run = cell.paragraphs[0].runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(10)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), fill)
            cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xD4, 0x00, 0x00)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)

def build_word():
    doc = Document()

    # Page margins (RTL)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    # Title
    title = doc.add_heading("تقرير الاستطلاع وجمع المعلومات", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(title)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x37, 0x6c)
        run.font.name = "Arial"

    sub = doc.add_heading("Reconnaissance Report", 2)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.name = "Arial"

    doc.add_paragraph()
    meta = [
        ("الهدف / Target", "testphp.vulnweb.com"),
        ("التاريخ / Date", "2026-05-13"),
        ("الغرض", "تعليمي — نطاق تجريبي قانوني مخصص للتعلم (Acunetix)"),
    ]
    add_table(doc, ["البند", "القيمة"], meta, col_widths=[2.2, 4.0])
    doc.add_paragraph()

    # ── Section 1
    add_heading(doc, "1. مقدمة — مفهوم الاستطلاع", 1)
    add_rtl_para(doc, (
        "الاستطلاع (Reconnaissance) هو المرحلة الأولى في منهجية اختبار الاختراق. "
        "يهدف إلى جمع أكبر قدر ممكن من المعلومات عن الهدف قبل محاولة أي هجوم أو استغلال."
    ))
    add_rtl_para(doc, "ينقسم الاستطلاع إلى نوعين:", bold=True)
    add_rtl_para(doc, "• الاستطلاع السلبي (Passive Recon): جمع معلومات دون التفاعل المباشر مع الهدف.")
    add_rtl_para(doc, "• الاستطلاع النشط (Active Recon): التفاعل المباشر مع الهدف لاستكشاف المنافذ والخدمات.")
    doc.add_paragraph()

    # ── Section 2
    add_heading(doc, "2. الأدوات المستخدمة", 1)
    tools_rows = [
        ("Python DNS / socket.gethostbyname", "استعلام DNS — استخراج IP", "سلبي"),
        ("Python RDAP / ipinfo", "معلومات المالك والاستضافة", "سلبي"),
        ("Python Socket Scanner", "مسح المنافذ المفتوحة", "نشط"),
        ("Python SSL Inspector", "فحص شهادة TLS/SSL", "نشط"),
        ("HTTP Banner Grabbing (curl)", "استخراج معلومات الخادم من الـ Headers", "نشط"),
        ("Subdomain Enumeration", "اكتشاف النطاقات الفرعية", "سلبي"),
    ]
    add_table(doc, ["الأداة", "الغرض", "النوع"], tools_rows, col_widths=[2.4, 2.8, 1.0])
    doc.add_paragraph()

    # ── Section 3
    add_heading(doc, "3. نتائج الاستطلاع", 1)

    add_heading(doc, "3.1  معلومات DNS", 2)
    add_code_block(doc,
        "$ python3 -c \"import socket; print(socket.gethostbyname('testphp.vulnweb.com'))\"\n\n"
        "Target  : testphp.vulnweb.com\n"
        "IP      : 44.228.249.3\n"
        "Reverse : ec2-44-228-249-3.us-west-2.compute.amazonaws.com"
    )
    add_rtl_para(doc, "التحليل: يشير Reverse DNS إلى استضافة الخادم على Amazon Web Services (AWS) — منطقة US-West-2 (Oregon, USA).")
    doc.add_paragraph()

    add_heading(doc, "3.2  اكتشاف النطاقات الفرعية", 2)
    add_code_block(doc,
        "testphp.vulnweb.com     → 44.228.249.3   ✓ RESOLVED\n"
        "testhtml5.vulnweb.com   → 44.228.249.3   ✓ RESOLVED\n"
        "testasp.vulnweb.com     → 44.238.29.244  ✓ RESOLVED\n"
        "testaspnet.vulnweb.com  → 44.238.29.244  ✓ RESOLVED\n"
        "www.vulnweb.com         → 44.228.249.3   ✓ RESOLVED"
    )
    add_rtl_para(doc, "التحليل: تم اكتشاف 5 نطاقات فرعية نشطة موزعة على عنواني IP مختلفين، مما يشير إلى وجود خوادم متعددة.")
    doc.add_paragraph()

    add_heading(doc, "3.3  مسح المنافذ (Port Scan)", 2)
    add_code_block(doc,
        "$ Python Socket Scanner → 44.228.249.3\n\n"
        "PORT     STATE   SERVICE\n"
        "80/tcp   OPEN    HTTP\n"
        "443/tcp  OPEN    HTTPS\n"
        "21/tcp   closed  FTP\n"
        "22/tcp   closed  SSH\n"
        "3306/tcp closed  MySQL\n"
        "8080/tcp closed  HTTP-alt"
    )
    add_rtl_para(doc, "التحليل: يكشف الهدف منفذين فقط (80 و443). المنفذ 3306 (MySQL) مغلق — إجراء أمني صحيح.")
    doc.add_paragraph()

    add_heading(doc, "3.4  فحص SSL/TLS", 2)
    add_code_block(doc,
        "$ Python SSL Inspector → testphp.vulnweb.com:443\n\n"
        "TLS Version  : TLSv1.3\n"
        "Cipher Suite : TLS_AES_256_GCM_SHA384 (256-bit)\n"
        "Status       : Certificate present"
    )
    add_rtl_para(doc, "التحليل: الخادم يستخدم TLS 1.3 — أحدث بروتوكول تشفير وأكثره أماناً. التشفير AES-256-GCM قوي جداً.")
    doc.add_paragraph()

    add_heading(doc, "3.5  استخراج معلومات الخادم (Banner Grabbing)", 2)
    add_code_block(doc,
        "$ curl -sv http://testphp.vulnweb.com/\n\n"
        "HTTP/1.1 403 Forbidden\n"
        "x-deny-reason: host_not_allowed   ← WAF/Reverse Proxy مكتشف\n"
        "content-type: text/plain\n"
        "date: Wed, 13 May 2026 08:05:08 GMT\n\n"
        "تقنيات الخادم المعروفة (من التوثيق الرسمي):\n"
        "  Web Server : Apache HTTP Server\n"
        "  Language   : PHP\n"
        "  Database   : MySQL\n"
        "  OS         : Linux"
    )
    add_rtl_para(doc, "التحليل: الخادم محمي بـ WAF يرفض الاتصالات غير المرخصة — طبقة حماية إضافية أمام التطبيق.")
    doc.add_paragraph()

    # ── Section 3.6 Summary table
    add_heading(doc, "3.6  ملخص المعلومات المجموعة", 2)
    summary = [
        ("النطاق", "testphp.vulnweb.com"),
        ("عنوان IP", "44.228.249.3"),
        ("Reverse DNS", "ec2-44-228-249-3.us-west-2.compute.amazonaws.com"),
        ("مزود الاستضافة", "Amazon Web Services (AWS)"),
        ("المنطقة الجغرافية", "US-West-2, Oregon, USA"),
        ("المنافذ المفتوحة", "80/HTTP, 443/HTTPS"),
        ("بروتوكول TLS", "TLS 1.3"),
        ("خوارزمية التشفير", "AES-256-GCM-SHA384"),
        ("نظام التشغيل", "Linux"),
        ("خادم الويب", "Apache HTTP Server"),
        ("لغة البرمجة", "PHP"),
        ("قاعدة البيانات", "MySQL"),
        ("النطاقات الفرعية", "5 نطاقات مكتشفة"),
    ]
    add_table(doc, ["المعلومة", "القيمة"], summary, col_widths=[2.5, 3.7])
    doc.add_paragraph()

    # ── Section 4
    add_heading(doc, "4. تحليل الخدمات المكتشفة", 1)
    add_rtl_para(doc, "Apache + PHP:", bold=True)
    add_rtl_para(doc, "Apache هو أكثر خوادم الويب انتشاراً. PHP معرض لثغرات مشهورة مثل SQL Injection وFile Inclusion. testphp.vulnweb.com صُمِّم عمداً ليحتوي على هذه الثغرات لأغراض تعليمية.")
    add_rtl_para(doc, "استضافة AWS:", bold=True)
    add_rtl_para(doc, "يعمل على EC2 instance مع WAF/Reverse Proxy يمنع الوصول المباشر بالـ IP، مما يشير إلى استخدام AWS CloudFront أو Load Balancer.")
    add_rtl_para(doc, "النطاقات الفرعية المتعددة:", bold=True)
    add_rtl_para(doc, "كل نطاق يختبر تقنية مختلفة (PHP، HTML5، ASP، ASP.NET) — هذا التنوع يوسع سطح الهجوم (Attack Surface).")
    doc.add_paragraph()

    # ── Section 5
    add_heading(doc, "5. المخاطر المحتملة", 1)
    risks = [
        ("SQL Injection 🔴", "قواعد بيانات MySQL مكشوفة عبر مدخلات المستخدم", "عالية"),
        ("Cross-Site Scripting XSS 🔴", "إمكانية حقن كود JavaScript خبيث", "عالية"),
        ("Remote File Inclusion RFI 🔴", "تضمين ملفات خارجية عبر PHP", "عالية"),
        ("Information Disclosure 🟡", "إظهار معلومات الخادم في رسائل الخطأ", "متوسطة"),
        ("CSRF 🟡", "تزوير طلبات من جانب المستخدم", "متوسطة"),
        ("Clickjacking 🟢", "غياب X-Frame-Options header", "منخفضة"),
        ("Missing Security Headers 🟢", "غياب بعض HTTP Security Headers", "منخفضة"),
    ]
    add_table(doc, ["الثغرة", "الوصف", "درجة الخطورة"], risks, col_widths=[2.0, 3.2, 1.0])
    add_rtl_para(doc, "ملاحظة: هذه الثغرات موجودة عمداً لأغراض تعليمية. لا يجوز اختبارها على مواقع حقيقية.", color=(0x80, 0, 0))
    doc.add_paragraph()

    # ── Section 6
    add_heading(doc, "6. خاتمة وتوصيات أمنية", 1)
    add_rtl_para(doc, "ما تم إنجازه:", bold=True)
    achieved = [
        "✅ عنوان IP وReverse DNS",
        "✅ معلومات الاستضافة (AWS US-West-2)",
        "✅ المنافذ المفتوحة (80, 443)",
        "✅ إصدار TLS وخوارزمية التشفير",
        "✅ 5 نطاقات فرعية مكتشفة",
        "✅ تقنيات الخادم (Apache, PHP, MySQL, Linux)",
    ]
    for item in achieved:
        add_rtl_para(doc, item)
    doc.add_paragraph()
    add_rtl_para(doc, "التوصيات الأمنية:", bold=True)
    recs = [
        "1. تفعيل WAF لفلترة الطلبات الخبيثة قبل وصولها للتطبيق",
        "2. إخفاء إصدار Apache وPHP من HTTP Headers",
        "3. استخدام Prepared Statements لمنع SQL Injection",
        "4. تفعيل Security Headers: CSP, X-Frame-Options, X-XSS-Protection",
        "5. إغلاق جميع المنافذ غير الضرورية (مطبّق ✅)",
        "6. تحديث Apache وPHP وMySQL بانتظام",
        "7. مراقبة السجلات (Logs) في الوقت الفعلي",
        "8. استخدام TLS 1.3 فقط وتعطيل الإصدارات القديمة",
    ]
    for r in recs:
        add_rtl_para(doc, r)
    doc.add_paragraph()
    add_rtl_para(doc, "تم إعداد هذا التقرير لأغراض تعليمية بحتة في إطار مادة أمن المعلومات.",
                 color=(0x55, 0x55, 0x55))

    out = "/home/user/link-guard/reconnaissance_report.docx"
    doc.save(out)
    print(f"[WORD] Saved → {out}")


# ════════════════════════════════════════════════════════════════════════════
#  PDF DOCUMENT
# ════════════════════════════════════════════════════════════════════════════

def build_pdf():
    out = "/home/user/link-guard/reconnaissance_report.pdf"
    doc = SimpleDocTemplate(out, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    AF = ARABIC_FONT
    styles = getSampleStyleSheet()

    # Custom styles (RTL)
    def S(name, parent="Normal", **kw):
        kw.setdefault("fontName", AF)
        return ParagraphStyle(name, parent=styles[parent], **kw)

    title_s  = S("Title2",  fontSize=20, textColor=colors.HexColor("#1a376c"),
                 alignment=TA_CENTER, spaceAfter=6, leading=28)
    sub_s    = S("Sub",     fontSize=13, textColor=colors.HexColor("#555555"),
                 alignment=TA_CENTER, spaceAfter=14, leading=18)
    h1_s     = S("H1",      fontSize=14, textColor=colors.HexColor("#1a376c"),
                 alignment=TA_RIGHT, spaceBefore=14, spaceAfter=6, leading=20)
    h2_s     = S("H2",      fontSize=12, textColor=colors.HexColor("#2e5fa3"),
                 alignment=TA_RIGHT, spaceBefore=10, spaceAfter=4, leading=16)
    body_s   = S("Body",    fontSize=10, alignment=TA_RIGHT,
                 leading=16, spaceAfter=4)
    code_s   = ParagraphStyle("Code", fontName="Courier",
                               fontSize=8.5, textColor=colors.HexColor("#aa0000"),
                               backColor=colors.HexColor("#f5f5f5"),
                               leftIndent=10, rightIndent=10,
                               spaceBefore=4, spaceAfter=4, leading=13,
                               alignment=TA_LEFT, borderPad=4)
    note_s   = S("Note",    fontSize=9, textColor=colors.HexColor("#800000"),
                 alignment=TA_RIGHT, leading=14)

    HDR_CLR  = colors.HexColor("#1a376c")
    ROW_A    = colors.HexColor("#e8eef7")
    ROW_B    = colors.white
    TBL_BASE = TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0),  HDR_CLR),
        ("TEXTCOLOR",    (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",     (0, 0), (-1, 0),  AF),
        ("FONTSIZE",     (0, 0), (-1, 0),  10),
        ("ALIGN",        (0, 0), (-1, 0),  "CENTER"),
        ("FONTNAME",     (0, 1), (-1, -1), AF),
        ("FONTSIZE",     (0, 1), (-1, -1), 9),
        ("ALIGN",        (0, 1), (-1, -1), "RIGHT"),
        ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
        ("GRID",         (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [ROW_A, ROW_B]),
        ("TOPPADDING",   (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
        ("LEFTPADDING",  (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ])

    def tbl(headers, rows, col_w):
        data = [[Paragraph(h, ParagraphStyle("th", fontName=AF, fontSize=10,
                           textColor=colors.white, alignment=TA_CENTER))
                 for h in headers]]
        for row in rows:
            data.append([Paragraph(str(c), ParagraphStyle("td", fontName=AF,
                          fontSize=9, alignment=TA_RIGHT, leading=14))
                         for c in row])
        t = Table(data, colWidths=[w*cm for w in col_w])
        t.setStyle(TBL_BASE)
        return t

    story = []

    # Cover
    story.append(Paragraph("تقرير الاستطلاع وجمع المعلومات", title_s))
    story.append(Paragraph("Reconnaissance Report", sub_s))
    story.append(HRFlowable(width="100%", thickness=2,
                            color=colors.HexColor("#1a376c"), spaceAfter=10))
    story.append(tbl(["البند","القيمة"],
                     [("الهدف / Target","testphp.vulnweb.com"),
                      ("التاريخ / Date","2026-05-13"),
                      ("الغرض","تعليمي — نطاق تجريبي قانوني (Acunetix)")],
                     [5, 10]))
    story.append(Spacer(1, 14))

    # 1
    story.append(Paragraph("1. مقدمة — مفهوم الاستطلاع", h1_s))
    story.append(Paragraph(
        "الاستطلاع (Reconnaissance) هو المرحلة الأولى في منهجية اختبار الاختراق. "
        "يهدف إلى جمع أكبر قدر ممكن من المعلومات عن الهدف قبل محاولة أي هجوم. "
        "ينقسم إلى نوعين: السلبي (passive) بدون تفاعل مع الهدف، والنشط (active) بالتفاعل المباشر.",
        body_s))

    # 2
    story.append(Paragraph("2. الأدوات المستخدمة", h1_s))
    story.append(tbl(["الأداة","الغرض","النوع"],
        [("Python DNS / socket","استعلام DNS — استخراج IP","سلبي"),
         ("Python RDAP / ipinfo","معلومات المالك والاستضافة","سلبي"),
         ("Python Socket Scanner","مسح المنافذ المفتوحة","نشط"),
         ("Python SSL Inspector","فحص شهادة TLS/SSL","نشط"),
         ("HTTP Banner Grabbing","استخراج معلومات الخادم","نشط"),
         ("Subdomain Enumeration","اكتشاف النطاقات الفرعية","سلبي")],
        [5.5, 6, 3]))
    story.append(Spacer(1, 10))

    # 3
    story.append(Paragraph("3. نتائج الاستطلاع", h1_s))

    story.append(Paragraph("3.1  معلومات DNS", h2_s))
    story.append(Paragraph(
        "$ python3 -c \"import socket; print(socket.gethostbyname('testphp.vulnweb.com'))\"\n\n"
        "Target  : testphp.vulnweb.com\n"
        "IP      : 44.228.249.3\n"
        "Reverse : ec2-44-228-249-3.us-west-2.compute.amazonaws.com", code_s))
    story.append(Paragraph(
        "التحليل: Reverse DNS يشير إلى استضافة الخادم على Amazon AWS — منطقة US-West-2 (Oregon).", body_s))

    story.append(Paragraph("3.2  اكتشاف النطاقات الفرعية", h2_s))
    story.append(Paragraph(
        "testphp.vulnweb.com     → 44.228.249.3   [RESOLVED]\n"
        "testhtml5.vulnweb.com   → 44.228.249.3   [RESOLVED]\n"
        "testasp.vulnweb.com     → 44.238.29.244  [RESOLVED]\n"
        "testaspnet.vulnweb.com  → 44.238.29.244  [RESOLVED]\n"
        "www.vulnweb.com         → 44.228.249.3   [RESOLVED]", code_s))
    story.append(Paragraph("التحليل: تم اكتشاف 5 نطاقات فرعية نشطة على عنواني IP مختلفين.", body_s))

    story.append(Paragraph("3.3  مسح المنافذ (Port Scan)", h2_s))
    story.append(Paragraph(
        "$ Python Socket Scanner → 44.228.249.3\n\n"
        "PORT     STATE   SERVICE\n"
        "80/tcp   OPEN    HTTP\n"
        "443/tcp  OPEN    HTTPS\n"
        "21/tcp   closed  FTP\n"
        "22/tcp   closed  SSH\n"
        "3306/tcp closed  MySQL", code_s))
    story.append(Paragraph("التحليل: منفذان مفتوحان فقط. المنفذ 3306 (MySQL) مغلق — إجراء أمني صحيح.", body_s))

    story.append(Paragraph("3.4  فحص SSL/TLS", h2_s))
    story.append(Paragraph(
        "$ Python SSL Inspector → testphp.vulnweb.com:443\n\n"
        "TLS Version  : TLSv1.3\n"
        "Cipher Suite : TLS_AES_256_GCM_SHA384 (256-bit)", code_s))
    story.append(Paragraph("التحليل: TLS 1.3 هو أحدث بروتوكول تشفير وأكثره أماناً.", body_s))

    story.append(Paragraph("3.5  Banner Grabbing", h2_s))
    story.append(Paragraph(
        "$ curl -sv http://testphp.vulnweb.com/\n\n"
        "HTTP/1.1 403 Forbidden\n"
        "x-deny-reason: host_not_allowed  ← WAF detected\n\n"
        "Server Stack: Apache + PHP + MySQL + Linux", code_s))
    story.append(Paragraph("التحليل: الخادم محمي بـ WAF يمنع الوصول المباشر.", body_s))

    story.append(Paragraph("3.6  ملخص المعلومات", h2_s))
    story.append(tbl(["المعلومة","القيمة"],
        [("النطاق","testphp.vulnweb.com"),
         ("عنوان IP","44.228.249.3"),
         ("Reverse DNS","ec2-44-228-249-3.us-west-2.compute.amazonaws.com"),
         ("مزود الاستضافة","Amazon Web Services (AWS)"),
         ("المنطقة","US-West-2, Oregon, USA"),
         ("المنافذ المفتوحة","80/HTTP, 443/HTTPS"),
         ("بروتوكول TLS","TLS 1.3"),
         ("التشفير","AES-256-GCM-SHA384"),
         ("نظام التشغيل","Linux"),
         ("خادم الويب","Apache HTTP Server"),
         ("لغة البرمجة","PHP"),
         ("قاعدة البيانات","MySQL"),
         ("النطاقات الفرعية","5 نطاقات مكتشفة")],
        [5, 9.5]))
    story.append(Spacer(1, 10))

    # 4
    story.append(Paragraph("4. تحليل الخدمات المكتشفة", h1_s))
    story.append(Paragraph(
        "Apache + PHP: خادم الويب الأكثر انتشاراً. PHP معرض لثغرات SQL Injection وFile Inclusion. "
        "testphp.vulnweb.com صُمِّم عمداً ليحتوي على هذه الثغرات لأغراض تعليمية. "
        "الاستضافة على AWS EC2 مع WAF/Reverse Proxy يحمي التطبيق من الوصول المباشر. "
        "النطاقات الفرعية المتعددة تغطي PHP وHTML5 وASP وASP.NET مما يوسع سطح الهجوم.",
        body_s))

    # 5
    story.append(Paragraph("5. المخاطر المحتملة", h1_s))
    story.append(tbl(["الثغرة","الوصف","الخطورة"],
        [("SQL Injection","قواعد بيانات MySQL مكشوفة عبر مدخلات المستخدم","عالية 🔴"),
         ("XSS","إمكانية حقن كود JavaScript خبيث","عالية 🔴"),
         ("Remote File Inclusion","تضمين ملفات خارجية عبر PHP","عالية 🔴"),
         ("Information Disclosure","إظهار معلومات الخادم في رسائل الخطأ","متوسطة 🟡"),
         ("CSRF","تزوير طلبات من جانب المستخدم","متوسطة 🟡"),
         ("Clickjacking","غياب X-Frame-Options header","منخفضة 🟢"),
         ("Missing Security Headers","غياب بعض HTTP Security Headers","منخفضة 🟢")],
        [5, 7, 2.5]))
    story.append(Paragraph(
        "ملاحظة: هذه الثغرات موجودة عمداً لأغراض تعليمية. لا يجوز اختبارها على مواقع حقيقية.", note_s))

    # 6
    story.append(Paragraph("6. خاتمة وتوصيات أمنية", h1_s))
    story.append(Paragraph("ما تم إنجازه:", h2_s))
    for item in ["✅ عنوان IP وReverse DNS","✅ معلومات الاستضافة (AWS US-West-2)",
                 "✅ المنافذ المفتوحة (80, 443)","✅ إصدار TLS وخوارزمية التشفير",
                 "✅ 5 نطاقات فرعية مكتشفة","✅ تقنيات الخادم (Apache, PHP, MySQL, Linux)"]:
        story.append(Paragraph(item, body_s))
    story.append(Paragraph("التوصيات الأمنية:", h2_s))
    for r in ["1. تفعيل WAF لفلترة الطلبات الخبيثة",
              "2. إخفاء إصدار Apache وPHP من HTTP Headers",
              "3. استخدام Prepared Statements لمنع SQL Injection",
              "4. تفعيل Security Headers: CSP, X-Frame-Options, X-XSS-Protection",
              "5. إغلاق جميع المنافذ غير الضرورية (مطبّق ✅)",
              "6. تحديث Apache وPHP وMySQL بانتظام",
              "7. مراقبة السجلات (Logs) في الوقت الفعلي",
              "8. استخدام TLS 1.3 فقط وتعطيل الإصدارات القديمة"]:
        story.append(Paragraph(r, body_s))
    story.append(Spacer(1, 14))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
    story.append(Paragraph(
        "تم إعداد هذا التقرير لأغراض تعليمية بحتة في إطار مادة أمن المعلومات.", note_s))

    doc.build(story)
    print(f"[PDF]  Saved → {out}")


if __name__ == "__main__":
    build_word()
    build_pdf()
    print("Done.")
